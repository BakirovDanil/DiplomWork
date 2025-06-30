from textwrap import wrap
import re
from docx import Document
from lxml import etree


def checking_code(code: str, base_rules: dict, chapter: str):
    """
    Функция, для проверки полученного кода по базе правил кодирования
    :param chapter:
    :param code:
    :param base_rules:
    :return:
    """
    decoder = wrap(code, base_rules[chapter]["Длина части кода"])
    answer = ''
    for position in decoder:
        if position in base_rules[chapter]:
            answer += position
        else:
            return f"{answer} не определен в базе кодов"
    return True

async def search_chapter_codes(base_rules: dict, document: Document) -> dict:
    """
    Функция, которая проверяет кодирование абзацев
    :param base_rules:
    :param document:
    :return:
    """
    chapter_status = {
        "Общее количество заголовков": 0,
        "Количество верно закодированных абзацев": 0,
        "Количество незакодированных абзацев": 0,
        "Заголовки с нарушениями кодирования": {

        }
    }
    number_paragraph = 1
    for paragraph in document.paragraphs:
        if paragraph.style.name in base_rules["chapter_rules"]["Допустимые стили"]:
            chapter_status["Общее количество заголовков"] += 1
            match = re.findall(fr"\b\d{{{base_rules["chapter_rules"]["Длина всего кода"]}}}\b", paragraph.text)
            if match:
                check_code_result = checking_code(match[0], base_rules, "chapter_rules")
                if check_code_result:
                    chapter_status["Количество верно закодированных абзацев"] += 1
                else:
                    chapter_status["Заголовки с нарушениями кодирования"][number_paragraph] = {
                        "Номер заголовка": number_paragraph,
                        "Название заголовка": paragraph.text,
                        "Ошибка": check_code_result,
                        "Заключение": "Код не соответствует кодировке"
                    }
                    chapter_status["Количество незакодированных абзацев"] += 1
            else:
                chapter_status["Заголовки с нарушениями кодирования"][number_paragraph] = {
                    "Номер заголовка": number_paragraph,
                    "Название заголовка": paragraph.text,
                    "Ошибка": "Код не обнаружен или является не полным (не равно 12 символам)",
                    "Заключение": "Код не соответствует кодировке"
                }
                chapter_status["Количество незакодированных абзацев"] += 1
        else:
            continue
        number_paragraph += 1
    return chapter_status


async def search_picture_codes(base_rules: dict, document: Document) -> dict:
    """
    Функция, которая проверяет кодирование рисунков
    :param base_rules:
    :param document:
    :return:
    """
    picture_status = {
        "Общее количество рисунков": 0,
        "Количество верно закодированных рисунков": 0,
        "Количество незакодированных рисунков": 0,
        "Рисунки с нарушением кодирования": {

        }
    }
    picture_number = 1
    for paragraph in document.paragraphs:
        if (paragraph.style.name in base_rules["picture_rules"]["Допустимые стили"]
            and re.findall(f"Рис", paragraph.text)):
            picture_status["Общее количество рисунков"] += 1
            match = re.findall(fr"\b\d{{{base_rules["picture_rules"]["Длина всего кода"]}}}\b", paragraph.text)
            if match:
                check_code_result = checking_code(match[0], base_rules, "picture_rules")
                if check_code_result:
                    picture_status["Количество верно закодированных рисунков"] += 1
                else:
                    picture_status["Рисунки с нарушением кодирования"][picture_number] = {
                        "Номер рисунка": picture_number,
                        "Название рисунка": paragraph.text,
                        "Ошибка": check_code_result,
                        "Заключение": "Код не соответствует кодировке"
                    }
                    picture_status["Количество незакодированных рисунков"] += 1
            else:
                picture_status["Рисунки с нарушением кодирования"][picture_number] = {
                    "Номер рисунка": picture_number,
                    "Название рисунка": paragraph.text,
                    "Ошибка": "Код не обнаружен или является не полным (не равно 15 символам)",
                    "Заключение": "Код не соответствует кодировке"
                }
                picture_status["Количество незакодированных рисунков"] += 1
        else:
            continue
        picture_number += 1

    # повторная проверка, но уже по текстовым полям (если такие есть)
    xml_str = document.part._element.xml
    root = etree.fromstring(xml_str.encode('utf-8'))
    ns = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'v': 'urn:schemas-microsoft-com:vml'
    }
    textboxes = root.findall('.//v:textbox', namespaces=ns)
    for textbox in textboxes:
        texts = textbox.findall('.//w:t', namespaces=ns)
        full_text = ''.join([t.text for t in texts if t.text])
        if re.findall(f"Рис", full_text):
            picture_status["Общее количество рисунков"] += 1
            match = re.findall(fr"\b\d{{{base_rules['picture_rules']['Длина всего кода']}}}\b", full_text)
            if match:
                check_code_result = checking_code(match[0], base_rules, "picture_rules")
                if check_code_result:
                    picture_status["Количество верно закодированных рисунков"] += 1
                else:
                    picture_status["Рисунки с нарушением кодирования"][picture_number] = {
                        "Номер рисунка": picture_number,
                        "Название рисунка": full_text,
                        "Ошибка": check_code_result,
                        "Заключение": "Код не соответствует кодировке"
                    }
                    picture_status["Количество незакодированных рисунков"] += 1
            else:
                picture_status["Рисунки с нарушением кодирования"][picture_number] = {
                    "Номер рисунка": picture_number,
                    "Название рисунка": full_text,
                    "Ошибка": "Код не обнаружен или является не полным (не равно 15 символам)",
                    "Заключение": "Код не соответствует кодировке"
                }
                picture_status["Количество незакодированных рисунков"] += 1
            picture_number += 1

    return picture_status

async def check_len_subsection(base_rules: dict, document: Document):
    len_status = {

    }
    lenght_paragraphs = len(document.paragraphs)
    subsection_index = 0
    i = 0

    while i < lenght_paragraphs:
        paragraph = document.paragraphs[i]
        if paragraph.style.name in ["Heading 2", "Заголовок 2"]:
            subsection_index += 1
            word_count = 0
            len_status[subsection_index] = {
                "Заголовок": document.paragraphs[i].text
            }
            i += 1
            while i < lenght_paragraphs and document.paragraphs[i].style.name not in ("Heading 2", "Заголовок 2"):
                text = document.paragraphs[i].text
                words = re.findall(r"\w+", text)
                word_count += len(words)
                i += 1
            len_status[subsection_index] ["Количество слов"] = word_count
        else:
            i += 1

    return len_status