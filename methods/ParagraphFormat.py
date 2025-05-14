import json

from docx import Document

from methods.dictWork import dictionary_output


async def checking_paragraphs_on_line_spacing(base_rules: dict, doc: Document) -> dict:
    """
    Функция для проверки всех абзацев на соответствие заданным шрифтам и прочим параметрам
    :param base_rules:
    :param doc:
    :return:
    """
    number_paragraph = 1
    paragraph_status = {
        "Общее количество абзацев": len(doc.paragraphs),
        "Количество абзацев, соответствующих правилам": 0,
        "Нетекстовые абзацы": 0,
        "Количество абзацев с нарушениями": 0,
        "Абзацы с нарушениями": {

        }
    }
    def error_paragraph_adding(para):
        """
        Добавление информации об абзаце с нарушениями форматирования
        """
        paragraph_status["Количество абзацев с нарушениями"] += 1
        paragraph_status["Абзацы с нарушениями"][number_paragraph] = {
            "Номер абзаца": number_paragraph,
            "Часть текста абзаца": para.text
        }

    paragraphs = doc.paragraphs
    for paragraph in paragraphs:
        # Проверяет, является ли абзац текстовым или же нет
        if not paragraph._p.xpath(".//w:t/text()"):
            paragraph_status["Нетекстовые абзацы"] += 1
            number_paragraph +=1
            continue
        if paragraph.style.name in base_rules:
            print(paragraph.style.name)
            paragraph_format = paragraph.paragraph_format
            if (paragraph_format.line_spacing is None and
                    paragraph_format.first_line_indent is None and
                    paragraph_format.space_after is None and
                    paragraph_format.space_before is None):
                paragraph_status["Количество абзацев, соответствующих правилам"] += 1
            else:
                error_paragraph_adding(paragraph)
        else:
            error_paragraph_adding(paragraph)
        number_paragraph += 1
    return paragraph_status


# def main(path: str):
#     try:
#         document = Document(path)
#         check_result = checking_paragraphs_on_line_spacing(open_base_rules("../rules/font_rules.json") ,document)
#         dictionary_output(check_result)
#         with open('../paragraph_result.json', 'w', encoding='utf-8') as f:
#             json.dump(check_result, f, ensure_ascii = False, indent = 1)
#     except Exception as e:
#         print(f"Ошибка: {e}")
#
# if __name__ == "__main__":
#     main("C:\\Users\\danil\\Desktop\\Lectures\\DiplomWork\\ВКР, текст.docx")

